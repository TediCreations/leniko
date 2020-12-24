from docx import Document
from docx.shared import Cm

from products.models import Product


def export2WholeSaleCatalog(filepath):
	document = Document()

	objs = Product.objects.all()
	print(f"Products: {len(objs)}")


	document.add_heading("Leniko Jewelry", 0)
	p = document.add_paragraph('Wholesale catalog 2020')
	document.add_page_break()

	# Table
	table = document.add_table(rows=1, cols=3)
	hdr_cells = table.rows[0].cells
	hdr_cells[0].text = 'Title'
	hdr_cells[1].text = 'Photo'
	hdr_cells[2].text = 'SKU'

	for obj in objs:
		print(f"Generating '{obj}'")
		row_cells = table.add_row().cells
		row_cells[0].text = obj.getTitle()
		#row_cells[1].text = ""
		row_cells[2].text = str(obj.sku)

		#import os
		from leniko.settings import BASE_DIR
		try:
			p = BASE_DIR + obj.getPhoto().url
		except Exception:
			p = BASE_DIR + "/static/img/shop/placeholder.jpg"

		paragraph = row_cells[1].paragraphs[0]
		run = paragraph.add_run()
		try:
			run.add_picture(p, width = Cm(4))
		except Exception:
			row_cells[1].text = ""

	document.add_page_break()
	document.save(filepath)
	print("Done")
