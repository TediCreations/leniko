from docx import Document
from docx.shared import Cm

from django.conf import settings
from products.models import Product


def export2WholeSaleCatalog(filepath):
	document = Document()

	objs = Product.objects.all()
	print(f"Products: {len(objs)}")


	document.add_heading("Leniko Jewelry", 0)
	p = document.add_paragraph('Wholesale catalog 2020')
	document.add_page_break()

	# Table
	table = document.add_table(rows=1, cols=4)
	hdr_cells = table.rows[0].cells
	hdr_cells[0].text = 'Title'
	hdr_cells[1].text = 'Photo'
	hdr_cells[2].text = 'SKU'
	hdr_cells[3].text = 'Price(Euro)'

	for obj in objs:
		print(f"Generating '{obj}'")
		row_cells = table.add_row().cells
		row_cells[0].text = obj.getTitle()
		#row_cells[1].text = ""
		row_cells[2].text = str(obj.sku)
		row_cells[3].text = obj.getPrice()

		paragraph = row_cells[1].paragraphs[0]
		run = paragraph.add_run()

		photoUrl = str(obj.getPhoto().photo.url)
		p = settings.BASE_DIR + photoUrl

		try:
			run.add_picture(p, width = Cm(4))
		except Exception as e:
			print(p)
			print(e)

	document.add_page_break()
	document.save(filepath)
	print("Done")
